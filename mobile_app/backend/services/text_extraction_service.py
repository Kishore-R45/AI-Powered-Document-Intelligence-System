"""
Enterprise-grade text extraction service for documents.
Handles PDF (text-based AND scanned/image-based) and image OCR extraction.

Pipeline:
1. Detect file type from magic bytes (never trust client MIME type alone)
2. Try native PDF text extraction (PyPDF)
3. If text < threshold → automatically run OCR (pytesseract + pdf2image)
4. Clean extracted text (remove noise, normalize, preserve important data)
5. Validate extraction quality before returning

Supported formats:
- PDF (native text + OCR fallback for scanned)
- JPEG / PNG / WEBP / TIFF / BMP (OCR)
- DOCX (python-docx)
"""

from typing import Optional, Tuple
import io
import re
import os


class TextExtractionService:
    """Service for extracting text from documents with OCR fallback."""
    
    # Minimum text length to consider PDF as text-based (not scanned)
    MIN_TEXT_LENGTH = 50
    
    # ─── Magic byte signatures for file type detection ───
    # This is the MOST RELIABLE way to identify file types
    # (never trust client-sent MIME types)
    FILE_SIGNATURES = {
        b'%PDF':                          'application/pdf',
        b'\xff\xd8\xff':                  'image/jpeg',
        b'\x89PNG\r\n\x1a\n':            'image/png',
        b'RIFF':                          'image/webp',      # WebP starts with RIFF...WEBP
        b'II\x2a\x00':                    'image/tiff',      # TIFF little-endian
        b'MM\x00\x2a':                    'image/tiff',      # TIFF big-endian
        b'BM':                            'image/bmp',
        b'PK\x03\x04':                    'application/docx', # DOCX is a ZIP
    }
    
    # Extension → MIME type mapping (fallback)
    EXTENSION_MAP = {
        '.pdf':  'application/pdf',
        '.jpg':  'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png':  'image/png',
        '.webp': 'image/webp',
        '.tiff': 'image/tiff',
        '.tif':  'image/tiff',
        '.bmp':  'image/bmp',
        '.docx': 'application/docx',
        '.doc':  'application/msword',
    }
    
    @classmethod
    def detect_file_type(cls, file_bytes: bytes, filename: str = '', client_mime: str = '') -> str:
        """
        Detect the actual file type using a 3-tier strategy:
        
        1. Magic bytes (most reliable - reads actual file header)
        2. File extension (reliable if filename is available)
        3. Client-provided MIME type (least reliable, last resort)
        
        This ensures correct detection even when the client sends
        'application/octet-stream' or an incorrect MIME type.
        
        Returns:
            Detected MIME type string (e.g., 'application/pdf')
        """
        detected = ''
        
        # ── Strategy 1: Magic bytes (highest priority) ──
        if file_bytes and len(file_bytes) >= 8:
            header = file_bytes[:16]  # Read first 16 bytes
            
            for signature, mime in cls.FILE_SIGNATURES.items():
                if header[:len(signature)] == signature:
                    # Special case: WebP needs additional check (RIFF....WEBP)
                    if signature == b'RIFF' and len(file_bytes) >= 12:
                        if file_bytes[8:12] != b'WEBP':
                            continue
                    detected = mime
                    print(f"[TextExtraction] File type detected from magic bytes: {detected}")
                    break
        
        # ── Strategy 2: File extension (fallback) ──
        if not detected and filename:
            ext = os.path.splitext(filename)[1].lower()
            if ext in cls.EXTENSION_MAP:
                detected = cls.EXTENSION_MAP[ext]
                print(f"[TextExtraction] File type detected from extension '{ext}': {detected}")
        
        # ── Strategy 3: Client MIME type (last resort) ──
        if not detected and client_mime and client_mime != 'application/octet-stream':
            detected = client_mime
            print(f"[TextExtraction] Using client-provided MIME type: {detected}")
        
        # ── Final fallback ──
        if not detected:
            detected = 'application/octet-stream'
            print(f"[TextExtraction] WARNING: Could not detect file type, defaulting to {detected}")
        
        return detected
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from a PDF file using native text extraction.
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"  Warning: Could not extract text from page {page_num + 1}: {e}")
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return ""
    
    @staticmethod
    def run_ocr_on_pdf(file_bytes: bytes) -> str:
        """
        Run OCR on a PDF by converting pages to images first.
        Used for scanned/image-based PDFs.
        
        Flow:
        1. Convert PDF pages to images using pdf2image
        2. Preprocess each image (grayscale + threshold)
        3. Extract text using pytesseract
        4. Combine results
        """
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            from PIL import Image, ImageFilter, ImageEnhance
            
            print("  Running OCR on PDF (scanned document detected)...")
            
            # Convert PDF pages to images
            try:
                images = convert_from_bytes(file_bytes, dpi=300)
            except Exception as e:
                print(f"  pdf2image conversion failed: {e}")
                try:
                    images = convert_from_bytes(file_bytes, dpi=200)
                except Exception as e2:
                    print(f"  pdf2image fallback also failed: {e2}")
                    return ""
            
            text_parts = []
            
            for i, image in enumerate(images):
                try:
                    processed_image = TextExtractionService._preprocess_image(image)
                    page_text = pytesseract.image_to_string(
                        processed_image,
                        config='--oem 3 --psm 6'
                    )
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                        print(f"  OCR page {i + 1}: extracted {len(page_text)} chars")
                except Exception as e:
                    print(f"  OCR error on page {i + 1}: {e}")
                    continue
            
            result = '\n\n'.join(text_parts)
            print(f"  OCR total: extracted {len(result)} characters from {len(images)} pages")
            return result
            
        except ImportError as e:
            print(f"  OCR dependencies not available: {e}")
            print("  Install: pip install pdf2image pytesseract Pillow")
            return ""
        except Exception as e:
            print(f"  OCR processing error: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    @staticmethod
    def _preprocess_image(image):
        """
        Preprocess image for better OCR accuracy.
        Convert to grayscale, enhance contrast/sharpness, apply thresholding.
        """
        from PIL import Image, ImageFilter, ImageEnhance, ImageOps
        
        gray = image.convert('L')
        
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(2.0)
        
        enhancer = ImageEnhance.Sharpness(enhanced)
        sharp = enhancer.enhance(1.5)
        
        blurred = sharp.filter(ImageFilter.MedianFilter(size=3))
        
        threshold = 140
        binary = blurred.point(lambda x: 255 if x > threshold else 0, '1')
        
        return binary
    
    @staticmethod
    def extract_from_image(file_bytes: bytes) -> str:
        """
        Extract text from an image using OCR with preprocessing.
        """
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(io.BytesIO(file_bytes))
            processed = TextExtractionService._preprocess_image(image)
            text = pytesseract.image_to_string(
                processed,
                config='--oem 3 --psm 6'
            )
            return text.strip()
        except ImportError as e:
            print(f"OCR dependencies not available: {e}")
            return ""
        except Exception as e:
            print(f"OCR extraction error: {e}")
            return ""
    
    @classmethod
    def extract_text(cls, file_bytes: bytes, mime_type: str = '', filename: str = '') -> str:
        """
        Enterprise-grade text extraction from any supported document.
        
        Uses magic-byte detection to identify file type (never trusts MIME alone).
        For PDFs: tries native extraction first, falls back to OCR if text is too short.
        For images: runs OCR with preprocessing.
        For DOCX: extracts text from paragraphs and tables.
        
        Args:
            file_bytes: Raw file content as bytes
            mime_type: Client-provided MIME type (may be wrong, used as fallback only)
            filename: Original filename (used for extension-based detection)
            
        Returns:
            Extracted and cleaned text string
        """
        if not file_bytes:
            print("[TextExtraction] ERROR: Empty file bytes provided")
            return ""
        
        # ── Step 0: Detect actual file type (NEVER trust client MIME alone) ──
        detected_type = cls.detect_file_type(file_bytes, filename, mime_type)
        print(f"[TextExtraction] Processing file: {filename or 'unknown'}")
        print(f"[TextExtraction]   Client MIME: {mime_type}")
        print(f"[TextExtraction]   Detected type: {detected_type}")
        print(f"[TextExtraction]   File size: {len(file_bytes)} bytes")
        
        text = ""
        
        # ── PDF extraction (native + OCR fallback) ──
        if detected_type == 'application/pdf':
            # Strategy 1: Native PDF text extraction
            text = cls.extract_from_pdf(file_bytes)
            print(f"[TextExtraction] Native PDF extraction: {len(text)} characters")
            
            # Strategy 2: If text too short → scanned PDF → run OCR
            if len(text.strip()) < cls.MIN_TEXT_LENGTH:
                print(f"[TextExtraction] Text too short ({len(text.strip())} < {cls.MIN_TEXT_LENGTH}), attempting OCR...")
                ocr_text = cls.run_ocr_on_pdf(file_bytes)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    print(f"[TextExtraction] Using OCR text: {len(text)} characters")
                else:
                    print(f"[TextExtraction] OCR didn't improve results, keeping native text")
            
            # Strategy 3: If still too short, try with higher DPI OCR
            if len(text.strip()) < cls.MIN_TEXT_LENGTH:
                print(f"[TextExtraction] Attempting high-DPI OCR as last resort...")
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    images = convert_from_bytes(file_bytes, dpi=400)
                    high_dpi_parts = []
                    for img in images:
                        processed = cls._preprocess_image(img)
                        page_text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3')
                        if page_text and page_text.strip():
                            high_dpi_parts.append(page_text.strip())
                    high_dpi_text = '\n\n'.join(high_dpi_parts)
                    if len(high_dpi_text.strip()) > len(text.strip()):
                        text = high_dpi_text
                        print(f"[TextExtraction] High-DPI OCR succeeded: {len(text)} characters")
                except Exception as e:
                    print(f"[TextExtraction] High-DPI OCR failed: {e}")
        
        # ── Image extraction (OCR) ──
        elif detected_type.startswith('image/'):
            text = cls.extract_from_image(file_bytes)
            print(f"[TextExtraction] Image OCR extraction: {len(text)} characters")
        
        # ── DOCX extraction ──
        elif detected_type == 'application/docx':
            text = cls._extract_from_docx(file_bytes)
            print(f"[TextExtraction] DOCX extraction: {len(text)} characters")
        
        # ── Unknown type: attempt PDF extraction anyway (common for mistyped files) ──
        else:
            print(f"[TextExtraction] Unknown type '{detected_type}', attempting PDF extraction as heuristic...")
            text = cls.extract_from_pdf(file_bytes)
            if len(text.strip()) < cls.MIN_TEXT_LENGTH:
                print(f"[TextExtraction] PDF heuristic failed, attempting image OCR...")
                text = cls.extract_from_image(file_bytes)
        
        # ── Clean the extracted text ──
        cleaned = cls.clean_text(text)
        print(f"[TextExtraction] Final cleaned text: {len(cleaned)} characters")
        
        if not cleaned:
            print(f"[TextExtraction] WARNING: No text could be extracted from this file!")
        
        return cleaned
    
    @staticmethod
    def _extract_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from a DOCX file (Microsoft Word).
        Extracts from paragraphs and tables.
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except ImportError:
            print("[TextExtraction] python-docx not installed. Install: pip install python-docx")
            return ""
        except Exception as e:
            print(f"[TextExtraction] DOCX extraction error: {e}")
            return ""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Professional text cleaning pipeline.
        
        Rules:
        - Remove control characters
        - Normalize unicode
        - Remove non-ASCII noise while preserving important characters
        - Normalize whitespace
        - Preserve lines with PAN, ID numbers, dates, GPA/SGPA
        """
        if not text:
            return ""
        
        # Remove control characters (except newlines and tabs)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize unicode
        try:
            import unicodedata
            text = unicodedata.normalize('NFKD', text)
        except:
            pass
        
        # Remove non-ASCII noise, keep important chars: / - : . , ( ) @ # & + = _ % ' "
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
        
        # Normalize whitespace (multiple spaces → single space)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Normalize newlines (multiple blank lines → double newline)
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Clean up each line
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)
            elif cleaned_lines and cleaned_lines[-1] != '':
                cleaned_lines.append('')
        
        return '\n'.join(cleaned_lines).strip()
